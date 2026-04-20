"""
utils/resume_parser.py
Extract raw text from uploaded PDF or DOCX resume files.
"""
from __future__ import annotations

import io
import logging
from pathlib import Path

logger = logging.getLogger(__name__)


def parse_resume(file_path: str | Path) -> str:
    """
    Extract plain text from a PDF or DOCX resume.

    Args:
        file_path: path to the uploaded file

    Returns:
        Extracted text (may contain newlines/whitespace)
    """
    file_path = Path(file_path)
    suffix = file_path.suffix.lower()

    if suffix == ".pdf":
        return _parse_pdf(file_path)
    elif suffix in (".docx", ".doc"):
        return _parse_docx(file_path)
    elif suffix in (".txt", ".md"):
        return file_path.read_text(encoding="utf-8", errors="ignore")
    else:
        raise ValueError(f"Unsupported file type: {suffix}")


def _parse_pdf(path: Path) -> str:
    """Extract text from PDF using PyMuPDF (fast) with pdfminer fallback."""
    text = _parse_pdf_pymupdf(path)
    if not text.strip():
        logger.info("PyMuPDF returned empty text, trying pdfminer...")
        text = _parse_pdf_pdfminer(path)
    return text


def _parse_pdf_pymupdf(path: Path) -> str:
    try:
        import fitz  # PyMuPDF
        doc = fitz.open(str(path))
        pages = [page.get_text() for page in doc]
        doc.close()
        return "\n".join(pages)
    except Exception as exc:
        logger.warning("PyMuPDF failed: %s", exc)
        return ""


def _parse_pdf_pdfminer(path: Path) -> str:
    try:
        from pdfminer.high_level import extract_text
        return extract_text(str(path))
    except Exception as exc:
        logger.warning("pdfminer failed: %s", exc)
        return ""


def _parse_docx(path: Path) -> str:
    try:
        from docx import Document
        doc = Document(str(path))
        paragraphs = [p.text for p in doc.paragraphs]
        # Also extract from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    paragraphs.append(cell.text)
        return "\n".join(p for p in paragraphs if p.strip())
    except Exception as exc:
        logger.warning("python-docx parse failed: %s", exc)
        return ""


def parse_resume_from_bytes(content: bytes, filename: str) -> str:
    """Parse resume directly from uploaded bytes (for Streamlit)."""
    import tempfile
    suffix = Path(filename).suffix
    with tempfile.NamedTemporaryFile(suffix=suffix, delete=False) as tmp:
        tmp.write(content)
        tmp_path = Path(tmp.name)
    try:
        return parse_resume(tmp_path)
    finally:
        tmp_path.unlink(missing_ok=True)
