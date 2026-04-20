"""
backend/pdf_generator.py
Renders the structured resume dict → PDF (ReportLab) and DOCX (python-docx).
Both outputs are ATS-safe: no tables/columns for the main content,
clean linear flow that ATS parsers can read left-to-right.
"""
from __future__ import annotations

import io
import logging
from pathlib import Path
from typing import Optional

from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT, TA_JUSTIFY
from reportlab.lib.pagesizes import LETTER
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import inch
from reportlab.platypus import (
    HRFlowable, Paragraph, SimpleDocTemplate, Spacer,
)

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger(__name__)


# ── Colour palette ────────────────────────────────────────────────────────────
PRIMARY = colors.HexColor("#1a365d")    # dark navy
ACCENT = colors.HexColor("#2b6cb0")     # medium blue
DARK_GREY = colors.HexColor("#2d3748")
MID_GREY = colors.HexColor("#718096")
RULE_COLOR = colors.HexColor("#bee3f8")


# ═══════════════════════════════════════════════════════════════════════════════
#  PDF GENERATOR
# ═══════════════════════════════════════════════════════════════════════════════
def _pdf_styles():
    styles = getSampleStyleSheet()

    name_style = ParagraphStyle(
        "Name", fontName="Helvetica-Bold", fontSize=22,
        textColor=PRIMARY, alignment=TA_CENTER, spaceAfter=2,
    )
    contact_style = ParagraphStyle(
        "Contact", fontName="Helvetica", fontSize=9,
        textColor=MID_GREY, alignment=TA_CENTER, spaceAfter=6,
    )
    section_header_style = ParagraphStyle(
        "SectionHeader", fontName="Helvetica-Bold", fontSize=11,
        textColor=PRIMARY, spaceBefore=10, spaceAfter=2,
        borderPadding=(0, 0, 2, 0),
    )
    job_title_style = ParagraphStyle(
        "JobTitle", fontName="Helvetica-Bold", fontSize=10,
        textColor=DARK_GREY, spaceBefore=6, spaceAfter=0,
    )
    company_style = ParagraphStyle(
        "Company", fontName="Helvetica-Oblique", fontSize=9,
        textColor=ACCENT, spaceAfter=2,
    )
    bullet_style = ParagraphStyle(
        "Bullet", fontName="Helvetica", fontSize=9,
        textColor=DARK_GREY, leftIndent=14, firstLineIndent=-10,
        spaceAfter=1, leading=13,
    )
    body_style = ParagraphStyle(
        "Body", fontName="Helvetica", fontSize=9,
        textColor=DARK_GREY, spaceAfter=3, leading=13, alignment=TA_JUSTIFY,
    )
    skill_style = ParagraphStyle(
        "Skill", fontName="Helvetica", fontSize=9,
        textColor=DARK_GREY, spaceAfter=2,
    )

    return {
        "name": name_style, "contact": contact_style,
        "section": section_header_style, "job_title": job_title_style,
        "company": company_style, "bullet": bullet_style,
        "body": body_style, "skill": skill_style,
    }


def _hr(story):
    story.append(HRFlowable(width="100%", thickness=1, color=RULE_COLOR, spaceAfter=4))


def _section(story, title: str, styles: dict):
    story.append(Spacer(1, 4))
    story.append(Paragraph(title.upper(), styles["section"]))
    _hr(story)


def generate_pdf(resume: dict, output_path: str | Path) -> Path:
    """Render resume dict → ATS-friendly PDF."""
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = SimpleDocTemplate(
        str(output_path),
        pagesize=LETTER,
        leftMargin=0.65 * inch,
        rightMargin=0.65 * inch,
        topMargin=0.6 * inch,
        bottomMargin=0.6 * inch,
    )
    styles = _pdf_styles()
    story = []

    # ── Header ────────────────────────────────────────────────────────────────
    story.append(Paragraph(resume.get("name", ""), styles["name"]))
    contact_parts = [
        p for p in [
            resume.get("email"), resume.get("phone"),
            resume.get("location"), resume.get("linkedin"),
            resume.get("github"),
        ] if p
    ]
    story.append(Paragraph(" | ".join(contact_parts), styles["contact"]))

    # ── Summary ───────────────────────────────────────────────────────────────
    if resume.get("summary"):
        _section(story, "Professional Summary", styles)
        story.append(Paragraph(resume["summary"], styles["body"]))

    # ── Skills ────────────────────────────────────────────────────────────────
    skills = resume.get("skills", {})
    primary = skills.get("primary", [])
    secondary = skills.get("secondary", [])
    if primary or secondary:
        _section(story, "Technical Skills", styles)
        if primary:
            story.append(Paragraph(
                f"<b>Core:</b> {', '.join(primary)}", styles["skill"]
            ))
        if secondary:
            story.append(Paragraph(
                f"<b>Additional:</b> {', '.join(secondary)}", styles["skill"]
            ))

    # ── Experience ────────────────────────────────────────────────────────────
    experience = resume.get("experience", [])
    if experience:
        _section(story, "Professional Experience", styles)
        for exp in experience:
            story.append(Paragraph(exp.get("title", ""), styles["job_title"]))
            story.append(Paragraph(
                f"{exp.get('company', '')}  ·  {exp.get('location', '')}  |  "
                f"{exp.get('start_date', '')} – {exp.get('end_date', '')}",
                styles["company"]
            ))
            for b in exp.get("bullets", []):
                story.append(Paragraph(f"• {b}", styles["bullet"]))

    # ── Projects ──────────────────────────────────────────────────────────────
    projects = resume.get("projects", [])
    if projects:
        _section(story, "Projects", styles)
        for proj in projects:
            name = proj.get("name", "")
            url = proj.get("url", "")
            header = f"<b>{name}</b>" + (f"  <i>({url})</i>" if url else "")
            story.append(Paragraph(header, styles["job_title"]))
            if proj.get("description"):
                story.append(Paragraph(proj["description"], styles["body"]))
            for b in proj.get("bullets", []):
                story.append(Paragraph(f"• {b}", styles["bullet"]))

    # ── Education ─────────────────────────────────────────────────────────────
    education = resume.get("education", [])
    if education:
        _section(story, "Education", styles)
        for edu in education:
            gpa = f"  |  GPA: {edu['gpa']}" if edu.get("gpa") else ""
            story.append(Paragraph(
                f"<b>{edu.get('degree', '')}</b>  –  {edu.get('institution', '')}  |  {edu.get('year', '')}{gpa}",
                styles["body"]
            ))

    # ── Certifications ────────────────────────────────────────────────────────
    certs = resume.get("certifications", [])
    if certs:
        _section(story, "Certifications", styles)
        for cert in certs:
            story.append(Paragraph(f"• {cert}", styles["bullet"]))

    doc.build(story)
    logger.info("PDF generated: %s", output_path)
    return output_path


# ═══════════════════════════════════════════════════════════════════════════════
#  DOCX GENERATOR
# ═══════════════════════════════════════════════════════════════════════════════
def _docx_add_section_heading(doc: Document, title: str):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10)
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    # Add bottom border to paragraph
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "BEE3F8")
    pBdr.append(bottom)
    pPr.append(pBdr)


def generate_docx(resume: dict, output_path: str | Path) -> Path:
    """Render resume dict → ATS-friendly DOCX."""
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.left_margin = Inches(0.75)
        section.right_margin = Inches(0.75)
        section.top_margin = Inches(0.65)
        section.bottom_margin = Inches(0.65)

    # Default font
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # ── Name ──────────────────────────────────────────────────────────────────
    name_p = doc.add_paragraph()
    name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name_run = name_p.add_run(resume.get("name", ""))
    name_run.bold = True
    name_run.font.size = Pt(20)
    name_run.font.color.rgb = RGBColor(0x1A, 0x36, 0x5D)

    # ── Contact ───────────────────────────────────────────────────────────────
    contact_parts = [
        p for p in [
            resume.get("email"), resume.get("phone"),
            resume.get("location"), resume.get("linkedin"),
            resume.get("github"),
        ] if p
    ]
    contact_p = doc.add_paragraph(" | ".join(contact_parts))
    contact_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact_p.runs[0].font.size = Pt(9)
    contact_p.runs[0].font.color.rgb = RGBColor(0x71, 0x80, 0x96)

    # ── Summary ───────────────────────────────────────────────────────────────
    if resume.get("summary"):
        _docx_add_section_heading(doc, "Professional Summary")
        p = doc.add_paragraph(resume["summary"])
        p.paragraph_format.space_after = Pt(4)

    # ── Skills ────────────────────────────────────────────────────────────────
    skills = resume.get("skills", {})
    primary = skills.get("primary", [])
    secondary = skills.get("secondary", [])
    if primary or secondary:
        _docx_add_section_heading(doc, "Technical Skills")
        if primary:
            p = doc.add_paragraph()
            p.add_run("Core: ").bold = True
            p.add_run(", ".join(primary))
            p.paragraph_format.space_after = Pt(2)
        if secondary:
            p = doc.add_paragraph()
            p.add_run("Additional: ").bold = True
            p.add_run(", ".join(secondary))
            p.paragraph_format.space_after = Pt(2)

    # ── Experience ────────────────────────────────────────────────────────────
    experience = resume.get("experience", [])
    if experience:
        _docx_add_section_heading(doc, "Professional Experience")
        for exp in experience:
            # Title
            title_p = doc.add_paragraph()
            title_p.paragraph_format.space_before = Pt(6)
            title_p.paragraph_format.space_after = Pt(0)
            title_run = title_p.add_run(exp.get("title", ""))
            title_run.bold = True
            title_run.font.size = Pt(10)

            # Company + dates
            co_p = doc.add_paragraph()
            co_p.paragraph_format.space_after = Pt(2)
            co_run = co_p.add_run(
                f"{exp.get('company', '')}  ·  {exp.get('location', '')}  "
                f"|  {exp.get('start_date', '')} – {exp.get('end_date', '')}"
            )
            co_run.italic = True
            co_run.font.size = Pt(9)
            co_run.font.color.rgb = RGBColor(0x2B, 0x6C, 0xB0)

            # Bullets
            for b in exp.get("bullets", []):
                bp = doc.add_paragraph(style="List Bullet")
                bp.text = b
                bp.paragraph_format.space_after = Pt(1)
                for run in bp.runs:
                    run.font.size = Pt(9)

    # ── Projects ──────────────────────────────────────────────────────────────
    projects = resume.get("projects", [])
    if projects:
        _docx_add_section_heading(doc, "Projects")
        for proj in projects:
            pp = doc.add_paragraph()
            pp.paragraph_format.space_before = Pt(4)
            run = pp.add_run(proj.get("name", ""))
            run.bold = True
            if proj.get("url"):
                pp.add_run(f"  ({proj['url']})")
            if proj.get("description"):
                doc.add_paragraph(proj["description"]).paragraph_format.space_after = Pt(1)
            for b in proj.get("bullets", []):
                bp = doc.add_paragraph(style="List Bullet")
                bp.text = b

    # ── Education ─────────────────────────────────────────────────────────────
    education = resume.get("education", [])
    if education:
        _docx_add_section_heading(doc, "Education")
        for edu in education:
            gpa = f"  |  GPA: {edu['gpa']}" if edu.get("gpa") else ""
            ep = doc.add_paragraph()
            ep.add_run(f"{edu.get('degree', '')}").bold = True
            ep.add_run(f"  –  {edu.get('institution', '')}  |  {edu.get('year', '')}{gpa}")

    # ── Certifications ────────────────────────────────────────────────────────
    certs = resume.get("certifications", [])
    if certs:
        _docx_add_section_heading(doc, "Certifications")
        for cert in certs:
            doc.add_paragraph(cert, style="List Bullet")

    doc.save(str(output_path))
    logger.info("DOCX generated: %s", output_path)
    return output_path


def generate_cover_letter_docx(
    cover_letter_text: str,
    candidate_name: str,
    job_title: str,
    company: str,
    output_path: str | Path,
) -> Path:
    """Save cover letter as DOCX."""
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    for section in doc.sections:
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)

    # Heading
    h = doc.add_heading(f"Cover Letter – {job_title} at {company}", level=2)
    h.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()

    # Body (preserve paragraph breaks)
    for para in cover_letter_text.split("\n\n"):
        if para.strip():
            doc.add_paragraph(para.strip())

    doc.save(str(output_path))
    return output_path
